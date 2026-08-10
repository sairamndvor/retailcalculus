import os
import re
import urllib.parse
import requests
from bs4 import BeautifulSoup, NavigableString
from PIL import Image
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# Constants
BASE_DIR = r"d:\Company Domains\Retailcalculus - Dashboard\retailcalculuscom"
OUTPUT_FILE = os.path.join(BASE_DIR, "Retail_Calculus_Documentation.docx")
TEMP_IMAGE_DIR = os.path.join(BASE_DIR, "temp_images")

# Setup temporary image folder
os.makedirs(TEMP_IMAGE_DIR, exist_ok=True)

# Color Scheme (Sleek slate / dark enterprise theme)
COLOR_PRIMARY = RGBColor(15, 23, 42)     # #0F172A (Deep Slate)
COLOR_SECONDARY = RGBColor(30, 41, 59)   # #1E293B (Dark Slate Blue)
COLOR_ACCENT = RGBColor(197, 160, 89)    # #C5A059 (Muted Gold)
COLOR_TEXT = RGBColor(51, 65, 85)        # #334155 (Slate Gray)
COLOR_LINK = RGBColor(30, 64, 175)       # #1E40AF (Blue)

# List of pages in the order they should appear in the document
PAGES = [
    {"file": "index.html", "title": "Homepage"},
    {"file": "platform.html", "title": "Platform Specifications"},
    {"file": "customer-stories.html", "title": "Customer Stories"},
    {"file": "book-demo.html", "title": "Book a Demo"},
    {"file": "contact.html", "title": "Contact Us"}
]

def clean_bookmark_name(href):
    """Sanitizes bookmark names for Word compatibility."""
    name = href.replace('#', '').replace('.html', '').replace('-', '_').replace('/', '_').replace('.', '_')
    name = re.sub(r'[^a-zA-Z0-9_]', '', name)
    if not name:
        name = "section_top"
    # Ensure it starts with a letter
    if name[0].isdigit():
        name = "b_" + name
    return name[:38]  # Word bookmarks limit is 40 chars

def add_bookmark(paragraph, name):
    """Adds a bookmark start and end inside a paragraph."""
    p = paragraph._p
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), name)
    p.insert(0, bookmark_start)
    
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')
    p.append(bookmark_end)

def add_hyperlink(paragraph, url, text, color="1E40AF", underline=True):
    """Helper to add an external hyperlink."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def add_local_hyperlink(paragraph, bookmark_name, text, color="1E40AF", underline=True):
    """Helper to add a local hyperlink linking to a bookmark."""
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    new_run.append(rPr)
    text_node = OxmlElement('w:t')
    text_node.text = text
    new_run.append(text_node)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def parse_inline_elements(paragraph, element):
    """Helper to parse HTML inline tags inside a paragraph and add appropriate runs."""
    for child in element.children:
        if isinstance(child, NavigableString):
            text = str(child)
            # Remove consecutive spaces/newlines but keep single space
            text = re.sub(r'\s+', ' ', text)
            if text and text != ' ':
                paragraph.add_run(text)
        elif child.name in ['strong', 'b']:
            run = paragraph.add_run(child.get_text())
            run.bold = True
        elif child.name in ['em', 'i']:
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'a':
            href = child.get('href', '')
            text = child.get_text()
            if not text:
                continue
            if href.startswith('#') or href.endswith('.html') or '.html#' in href:
                bookmark_name = clean_bookmark_name(href)
                add_local_hyperlink(paragraph, bookmark_name, text)
            else:
                add_hyperlink(paragraph, href, text)
        elif child.name == 'br':
            paragraph.add_run().add_break()
        elif child.name in ['span', 'code', 'small']:
            run = paragraph.add_run(child.get_text())
            if child.name == 'code':
                run.font.name = 'Courier New'

def process_image(src):
    """Downloads or converts website images to PNG format for Word insertion."""
    img_name = os.path.basename(src.split('?')[0])
    img_name_png = os.path.splitext(img_name)[0] + ".png"
    local_png_path = os.path.join(TEMP_IMAGE_DIR, img_name_png)

    # 1. Download if remote
    if src.startswith('http://') or src.startswith('https://'):
        try:
            r = requests.get(src, timeout=10)
            if r.status_code == 200:
                temp_raw = os.path.join(TEMP_IMAGE_DIR, img_name)
                with open(temp_raw, 'wb') as f:
                    f.write(r.content)
                img = Image.open(temp_raw)
                img.save(local_png_path, "PNG")
                return local_png_path
        except Exception as e:
            print(f"Error downloading image {src}: {e}")
            return None

    # 2. Convert if local
    else:
        # Resolve relative local path
        local_path = os.path.join(BASE_DIR, src.replace('/', os.sep))
        if os.path.exists(local_path):
            try:
                img = Image.open(local_path)
                img.save(local_png_path, "PNG")
                return local_png_path
            except Exception as e:
                print(f"Error converting local image {local_path}: {e}")
                return None
    return None

def add_placeholder_box(doc, title, description):
    """Creates a styled placeholder block when images/videos cannot be directly embedded."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set borders and shading using XML
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        '<w:tblBorders %s>'
        '<w:top w:val="single" w:sz="6" w:space="0" w:color="cbd5e1"/>'
        '<w:bottom w:val="single" w:sz="6" w:space="0" w:color="cbd5e1"/>'
        '<w:left w:val="single" w:sz="6" w:space="0" w:color="cbd5e1"/>'
        '<w:right w:val="single" w:sz="6" w:space="0" w:color="cbd5e1"/>'
        '</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders)
    
    cell = table.cell(0, 0)
    # Set cell shading (light gray background)
    shading_elm = parse_xml(r'<w:shd {} w:fill="f8fafc"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    
    cell.width = Inches(5.5)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p.add_run(f"[{title}]\n")
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = COLOR_SECONDARY
    
    run_desc = p.add_run(description)
    run_desc.italic = True
    run_desc.font.size = Pt(9.5)
    run_desc.font.color.rgb = COLOR_TEXT

def create_document():
    doc = Document()

    # 1. Page Margins Setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 2. Base Typography Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # Headings Setup
    for h_name, size, color in [
        ('Heading 1', 22, COLOR_PRIMARY),
        ('Heading 2', 16, COLOR_SECONDARY),
        ('Heading 3', 13, COLOR_SECONDARY),
        ('Heading 4', 11, COLOR_SECONDARY)
    ]:
        style = doc.styles[h_name]
        style.font.name = 'Calibri'
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        if h_name == 'Heading 1':
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(8)
        elif h_name == 'Heading 2':
            style.paragraph_format.space_before = Pt(16)
            style.paragraph_format.space_after = Pt(6)
        else:
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(4)

    # 3. Cover Page Creation
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(100)
    title_p.paragraph_format.space_after = Pt(12)
    
    title_run = title_p.add_run("RETAIL CALCULUS")
    title_run.bold = True
    title_run.font.size = Pt(32)
    title_run.font.color.rgb = COLOR_PRIMARY
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(36)
    sub_run = subtitle_p.add_run("Enterprise Cloud Software Suite for Furniture, Appliance, & Big-Ticket Retailers\nComplete Website and Product Documentation")
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(100, 116, 139) # Slate-500
    
    # Try adding Cover Image (logo)
    logo_path = process_image("assets/images/logo-black.webp")
    if logo_path:
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(80)
        p_logo.add_run().add_picture(logo_path, width=Inches(2.5))
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.paragraph_format.space_before = Pt(80)
    meta_run = meta_p.add_run("Version 1.0  |  Compiled from Local Source Assets\nCo-Founded by Edward L Rennemann & Jaikishen Koranchath\nhttps://retailcalculus.com")
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = RGBColor(148, 163, 184) # Slate-400
    
    doc.add_page_break()

    # 4. Table of Contents Page
    toc_heading = doc.add_paragraph()
    toc_heading.style = 'Heading 1'
    toc_heading.add_run("Table of Contents")
    add_bookmark(toc_heading, "toc_top")
    
    toc_desc = doc.add_paragraph("Select a page from the list below to navigate directly to its section:")
    toc_desc.paragraph_format.space_after = Pt(18)

    for page in PAGES:
        p_item = doc.add_paragraph(style='Normal')
        p_item.paragraph_format.left_indent = Inches(0.25)
        p_item.paragraph_format.space_after = Pt(4)
        
        # Local link to section bookmark
        bookmark_name = clean_bookmark_name(page['file'])
        add_local_hyperlink(p_item, bookmark_name, page['title'])
        
        # Add descriptive tag
        desc_run = p_item.add_run(f" — Web Content ({page['file']})")
        desc_run.font.color.rgb = RGBColor(100, 116, 139)
        desc_run.font.size = Pt(9.5)

    doc.add_page_break()

    # Enable different first page header/footer (hiding it on cover/TOC pages is handled by Word)
    # We will set a custom footer with page numbers
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    
    # Configure Header / Footer
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = header_p.add_run("Retail Calculus — Platform Documentation")
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = RGBColor(148, 163, 184)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = footer_p.add_run("Confidential  |  Page ")
    frun.font.size = Pt(8.5)
    frun.font.color.rgb = RGBColor(148, 163, 184)
    # Add page number field using XML
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    footer_p._p.append(fldSimple)

    # 5. Extract and format HTML content
    for idx, page in enumerate(PAGES):
        file_path = os.path.join(BASE_DIR, page['file'])
        if not os.path.exists(file_path):
            print(f"Skipping missing file: {file_path}")
            continue

        print(f"Processing {page['file']}...")
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')

        # Create Section Heading
        h1 = doc.add_paragraph()
        h1.style = 'Heading 1'
        h1.paragraph_format.space_before = Pt(24)
        h_run = h1.add_run(page['title'])
        
        # Add local bookmark for TOC
        bookmark_name = clean_bookmark_name(page['file'])
        add_bookmark(h1, bookmark_name)

        # Parse body elements (excluding nav, footer, script)
        body = soup.body
        if not body:
            continue

        # Extract primary content elements in order
        for child in body.descendants:
            if child.name is None:
                continue
            
            # Avoid traversing grandchildren multiple times
            if child.parent.name in ['nav', 'footer', 'script', 'style', 'iframe'] or child.parent.get('class') == ['mobile-menu-overlay'] or child.get('class') == ['mobile-menu-overlay']:
                continue
            if child.name in ['nav', 'footer', 'script', 'style']:
                continue

            # Process top-level headings, paragraphs, lists, and assets
            if child.name == 'h2':
                text = child.get_text().strip()
                if text:
                    p = doc.add_paragraph(style='Heading 2')
                    add_bookmark(p, clean_bookmark_name(text))
                    parse_inline_elements(p, child)

            elif child.name == 'h3':
                text = child.get_text().strip()
                if text:
                    p = doc.add_paragraph(style='Heading 3')
                    add_bookmark(p, clean_bookmark_name(text))
                    parse_inline_elements(p, child)

            elif child.name == 'h4':
                text = child.get_text().strip()
                if text:
                    p = doc.add_paragraph(style='Heading 4')
                    parse_inline_elements(p, child)

            elif child.name == 'p':
                # Avoid capturing parent tags already processed
                if child.parent.name in ['p', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    continue
                # Skip stats numbers which are processed separately or would be too segmented
                text = child.get_text().strip()
                if text:
                    # check if list paragraph
                    if child.parent.name == 'li':
                        continue
                    p = doc.add_paragraph(style='Normal')
                    parse_inline_elements(p, child)

            elif child.name == 'li':
                text = child.get_text().strip()
                if text:
                    p = doc.add_paragraph(style='List Bullet')
                    parse_inline_elements(p, child)

            elif child.name == 'img':
                src = child.get('src', '')
                alt = child.get('alt', 'Image')
                if src:
                    # Don't embed tiny logo files repeatedly in content
                    if 'logo-black' in src and idx > 0:
                        continue
                    img_path = process_image(src)
                    if img_path and os.path.exists(img_path):
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_before = Pt(8)
                        p_img.paragraph_format.space_after = Pt(4)
                        try:
                            p_img.add_run().add_picture(img_path, width=Inches(4.5))
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_cap.paragraph_format.space_after = Pt(12)
                            cap_run = p_cap.add_run(f"Figure: {alt if alt else 'Website Image Asset'}")
                            cap_run.font.size = Pt(9.0)
                            cap_run.font.italic = True
                            cap_run.font.color.rgb = RGBColor(100, 116, 139)
                        except Exception as e:
                            print(f"Could not insert image {img_path}: {e}")
                            add_placeholder_box(doc, "Image Placeholder", f"Asset Location: {src}\nDescription: {alt}")
                    else:
                        add_placeholder_box(doc, "Image Placeholder", f"Asset Location: {src}\nDescription: {alt}")

            elif child.name == 'iframe':
                # Video embeds (e.g. YouTube demo)
                src = child.get('src', '')
                title = child.get('title', 'Video Demonstration')
                add_placeholder_box(doc, "Video Demonstration Placeholder", f"Embed URL: {src}\nTitle: {title}")

        # Add page break between website sections
        if idx < len(PAGES) - 1:
            doc.add_page_break()

    # Save generated document
    doc.save(OUTPUT_FILE)
    print(f"Successfully generated documentation at {OUTPUT_FILE}")

if __name__ == "__main__":
    create_document()
